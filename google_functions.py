# -*- coding: utf-8 -*-

#Import required libraries
import pandas as pd
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from subprocess import Popen
from docx import Document
from time import sleep
import os.path
import shutil
import csv

def initiate():
    gauth = GoogleAuth()
    gauth.LocalWebserverAuth()
    drive = GoogleDrive(gauth)

    return(drive)

#Create set of new google docs per student for new project
def create_files_from_list(drive, class_name, project_name):
    
    #Get list of students and emails
    student_dict = pd.read_csv("".join(["resources/",class_name,"_grades.csv"]))
    student_dict = dict(zip(student_dict['user'], student_dict['email']))
    
    #Create new folder
    folder_metadata = {'title' : "_".join([class_name, project_name]),
                       'mimeType' : 'application/vnd.google-apps.folder'}
    folder = drive.CreateFile(folder_metadata)
    folder.Upload()
    
    #Add project to list
    with open('resources/projects.csv', 'r+') as f:
        reader = csv.reader(f)
        project_list = list(reader)
        project_list = [item for sublist in project_list for item in sublist]
    if not project_name in project_list:
        project_list.append(project_name)
        project_list = pd.Series(project_list)
        project_list.to_csv('resources/projects.csv', index=False)
    
    #Create name-link dictionary
    name_link = {}
    
    for student in student_dict.keys():
        
        try:
            #Set template
            template = "".join(["templates/", project_name, ".docx"])
            if not os.path.isfile(template):
                #Create word document
                document = Document()
                document.add_heading(project_name, level = 3)
                document.add_heading("Essay Title: ", level = 5)
                document.add_heading("Student Name: ", level = 5)
                
                #Save word document
                template = "".join(["templates/", project_name, ".docx"])
                document.save(template)
                
            #Make individualized files
            filename = "".join([str(student),"_",str(project_name),".docx"])
            shutil.copy(template, filename)
            
            #Upload to folder
            sleep(2)
            f = drive.CreateFile({'title': filename, "parents": [{'kind': 'drive#fileLink', 
                'id': folder['id']}], 'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
            f.SetContentFile(filename)
            f.Upload(param={'convert': True})
            sleep(2) #Avoid timeout
                                                    
            #Share file
            permission = f.InsertPermission({'type': 'user',
                                             'value': student_dict[student],
                                             'role': 'writer'})
            
            sleep(2)
            #Generate link
            name_link[student] = f['alternateLink']
            
            #Delete file
            print(filename)
            os.remove(filename)
            sleep(1) #Avoid timeout

        except Exception as e:
            print("Failed on %s" % (filename))
            print(e)

    #Confirmation
    print("All files created.")
    return(name_link)

#Returns the id(s) of specified files
def get_file_id(drive, filenames, parent = 'root'):
    _q = {'q': "'{}' in parents and trashed=false".format(parent)}
    file_list = drive.ListFile(_q).GetList()
    target_files = []
    for file in file_list:
        if filenames in file['title']:
            target_files.append(file)
    return(target_files)

#Download essays from google drive
def download_essays(drive, class_name, project_name, pre_post):
    #Select project folder
    folder = get_file_id(drive,"_".join([class_name, project_name]))[0]
    folder_id = folder['id']

    #Get list of files
    essay_list = get_file_id(drive, project_name, parent = folder_id)

    #Add file content to dataframe
    name = []
    essay_content = []
    word_count = []
    revision = []

    for essay in essay_list:
        name.append(essay['title'])
        try:
            content = essay.GetContentString(mimetype='text/plain', remove_bom=True)
        except(KeyError):
            content = "Error-Generated Content"
            print("".join(["Error with ", essay['title']]))
        
        #Fix curly quotes
        content = content.replace(u'\u201c', '\"').replace(u'\u201d', '\"')
        content = content.replace(u'\u2018', '\'').replace(u'\u2019', '\'')
        
        essay_content.append(content)
        word_count.append(len(content.split(' ')))
        revision.append(pre_post)

    #Save essays    
    essays = pd.DataFrame(data={'name': name, 
                                'essay_content': essay_content,
                                'revision': revision, 'word_choice': ['0']*len(name),
                                'word_count': word_count,
                                'transitional_phrases': ['0']*len(name), 
                                'sentence_length': ['0']*len(name),
                                'passive_voice': ['0']*len(name), 
                                'simple_starts': ['0']*len(name),
                                'vocabulary': ['0']*len(name),	
                                'grade': ['0']*len(name), 'letter': ['F']*len(name)})
    filename = "".join(["output/",str(class_name),"/",str(project_name),"/",str(project_name),".csv"])
    print(filename)
    
    #Create or Append
    if os.path.isfile(filename):
        old_essays = pd.read_csv(filename)
        essays = old_essays.append(essays, ignore_index=True)
    #Save
    essays.to_csv(filename, index=False)

    #Open saved document
    command = "".join(["open ",filename," -a \"Microsoft Excel\""])
    Popen(command, shell=True, stdin=None, stdout=None, stderr=None,
         close_fds=True, bufsize=0)

    print("Essays dowloaded.")
